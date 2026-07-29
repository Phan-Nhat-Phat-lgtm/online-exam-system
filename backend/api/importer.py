import re
import docx
import PyPDF2
import io
from .models import Question, QuestionBank

def parse_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Tách dòng nếu các dòng bị dính nhau trong một paragraph
            for sub_line in text.split('\n'):
                if sub_line.strip():
                    lines.append(sub_line.strip())
    # Check tables if content is inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    for sub_line in text.split('\n'):
                        if sub_line.strip():
                            lines.append(sub_line.strip())
    return lines

def parse_pdf(file_bytes):
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    lines = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            for line in text.split('\n'):
                line_str = line.strip()
                if line_str:
                    lines.append(line_str)
    return lines

def import_questions_from_file(file_obj, filename, bank_id):
    bank = QuestionBank.objects.get(id=bank_id)
    file_bytes = file_obj.read()
    
    if filename.lower().endswith('.docx'):
        lines = parse_docx(file_bytes)
    elif filename.lower().endswith('.pdf'):
        lines = parse_pdf(file_bytes)
    else:
        return {'success': False, 'error': 'Chỉ hỗ trợ file định dạng .docx hoặc .pdf'}

    if not lines:
        return {'success': False, 'error': 'File không chứa nội dung văn bản hợp lệ'}

    parsed_questions = []
    errors = []
    
    current_q = None
    
    # Regex patterns - Ultra-flexible for teacher's format
    q_pattern = re.compile(r'^(?:Câu|câu)\s*(\d+)[\:\.]?\s*(.*)', re.IGNORECASE)
    opt_a_pattern = re.compile(r'^[A][\.\)\/\:\s]\s*(.*)', re.IGNORECASE)
    opt_b_pattern = re.compile(r'^[B][\.\)\/\:\s]\s*(.*)', re.IGNORECASE)
    opt_c_pattern = re.compile(r'^[C][\.\)\/\:\s]\s*(.*)', re.IGNORECASE)
    opt_d_pattern = re.compile(r'^[D][\.\)\/\:\s]\s*(.*)', re.IGNORECASE)
    ans_pattern = re.compile(r'^(?:Đáp án|ĐÁP ÁN|Key|Đáp án đúng|Câu trả lời)\s*[\:\=\s]*([A-D])', re.IGNORECASE)

    for line_idx, line in enumerate(lines, start=1):
        line = line.strip()
        # Check if line matches a new Question
        q_match = q_pattern.match(line)
        ans_match = ans_pattern.search(line) # Dùng search thay vì match để linh hoạt hơn
        opt_a = opt_a_pattern.match(line)
        opt_b = opt_b_pattern.match(line)
        opt_c = opt_c_pattern.match(line)
        opt_d = opt_d_pattern.match(line)

        if ans_match:
            if current_q:
                current_q['correct_answer'] = ans_match.group(1).upper()
                # Validate current question before finalizing
                if not current_q['content']:
                    errors.append(f"Dòng {current_q['line_num']}: Câu hỏi thiếu nội dung")
                elif not current_q['option_a'] or not current_q['option_b'] or not current_q['option_c'] or not current_q['option_d']:
                    errors.append(f"Dòng {current_q['line_num']}: Câu hỏi '{current_q['content'][:30]}' thiếu các lựa chọn A, B, C, D")
                elif not current_q['correct_answer']:
                    errors.append(f"Dòng {current_q['line_num']}: Câu hỏi '{current_q['content'][:30]}' thiếu đáp án đúng")
                else:
                    parsed_questions.append(current_q)
                current_q = None
            else:
                errors.append(f"Dòng {line_idx}: Tìm thấy đáp án nhưng không thuộc câu hỏi nào: '{line}'")

        elif opt_a:
            if current_q:
                current_q['option_a'] = opt_a.group(1).strip()
            else:
                errors.append(f"Dòng {line_idx}: Đáp án A nằm ngoài câu hỏi: '{line}'")
        elif opt_b:
            if current_q:
                current_q['option_b'] = opt_b.group(1).strip()
            else:
                errors.append(f"Dòng {line_idx}: Đáp án B nằm ngoài câu hỏi: '{line}'")
        elif opt_c:
            if current_q:
                current_q['option_c'] = opt_c.group(1).strip()
            else:
                errors.append(f"Dòng {line_idx}: Đáp án C nằm ngoài câu hỏi: '{line}'")
        elif opt_d:
            if current_q:
                current_q['option_d'] = opt_d.group(1).strip()
            else:
                errors.append(f"Dòng {line_idx}: Đáp án D nằm ngoài câu hỏi: '{line}'")
        elif q_match:
            # Save previous unfinished question if exists
            if current_q:
                if current_q['content'] and current_q['option_a'] and current_q['option_b'] and current_q['option_c'] and current_q['option_d'] and current_q['correct_answer']:
                    parsed_questions.append(current_q)
                else:
                    errors.append(f"Dòng {current_q['line_num']}: Câu hỏi '{current_q['content'][:30]}' chưa đầy đủ thông tin")
                current_q = None
            
            content_text = q_match.group(2).strip()
            # If "Câu 1:" has no text on the same line, the next lines will be appended
            current_q = {
                'line_num': line_idx,
                'content': content_text,
                'option_a': '',
                'option_b': '',
                'option_c': '',
                'option_d': '',
                'correct_answer': '',
            }
        else:
            if current_q:
                if not current_q['option_a']:
                    # Multi-line question content continuation
                    current_q['content'] = (current_q['content'] + ' ' + line).strip()
                elif current_q['option_d']:
                    # Might be text after option D but before answer
                    pass
                elif current_q['option_c']:
                    current_q['option_c'] = (current_q['option_c'] + ' ' + line).strip()
                elif current_q['option_b']:
                    current_q['option_b'] = (current_q['option_b'] + ' ' + line).strip()
                elif current_q['option_a']:
                    current_q['option_a'] = (current_q['option_a'] + ' ' + line).strip()

    # Check unhandled last question if no explicit answer line found or finished
    if current_q:
        if current_q['content'] and current_q['option_a'] and current_q['option_b'] and current_q['option_c'] and current_q['option_d'] and current_q['correct_answer']:
            parsed_questions.append(current_q)
        else:
            errors.append(f"Dòng {current_q['line_num']}: Câu hỏi chưa đầy đủ thông tin (Nội dung, A, B, C, D, Đáp án)")

    if errors and not parsed_questions:
        return {
            'success': False,
            'errors': errors,
            'imported_count': 0,
            'message': 'Không thể import do có lỗi định dạng.'
        }

    # Save to Database
    created_objs = []
    for q_data in parsed_questions:
        obj = Question(
            bank=bank,
            content=q_data['content'],
            option_a=q_data['option_a'],
            option_b=q_data['option_b'],
            option_c=q_data['option_c'],
            option_d=q_data['option_d'],
            correct_answer=q_data['correct_answer']
        )
        created_objs.append(obj)

    Question.objects.bulk_create(created_objs)

    return {
        'success': True,
        'imported_count': len(created_objs),
        'errors': errors,
        'message': f"Đã import thành công {len(created_objs)} câu hỏi."
    }
